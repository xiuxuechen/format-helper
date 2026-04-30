"""生成 CODE-004 使用的标准 Word 样例和事实快照。"""

from __future__ import annotations

import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "fixtures" / "standard" / "official-report.docx"
SNAPSHOT_PATH = ROOT / "fixtures" / "standard" / "official-report.snapshot.json"
FIXED_CREATED_AT = "2026-04-30T16:30:00+00:00"
FIXED_ZIP_DATE = (2026, 4, 30, 16, 30, 0)


def set_run_font(run, east_asia: str, size_pt: float, bold: bool = False) -> None:
    """设置 run 的中英文字体、字号和加粗状态。"""
    run.font.name = east_asia
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5, first_line_cm: float | None = None) -> None:
    """设置段落行距、段前段后和首行缩进。"""
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_cm)


def add_paragraph(document: Document, text: str, style: str, font: str, size: float, bold: bool, first_line_cm: float | None = None):
    """新增段落并应用标准样例格式。"""
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, font, size, bold=bold)
    set_paragraph_spacing(paragraph, first_line_cm=first_line_cm)
    return paragraph


def create_docx() -> None:
    """创建正式材料标准 Word 样例。"""
    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("正式材料格式标准样例")
    set_run_font(run, "黑体", 22, bold=True)
    title.paragraph_format.space_after = Pt(18)

    add_paragraph(document, "一、总体要求", "Heading 1", "黑体", 16, True)
    add_paragraph(
        document,
        "本样例用于验证 format-helper v3 的语义规则草案与规则摘要生成链路。",
        "Normal",
        "宋体",
        12,
        False,
        first_line_cm=0.74,
    )
    add_paragraph(document, "（一）材料范围", "Heading 2", "楷体", 14, True)
    add_paragraph(
        document,
        "标准材料应保持标题层级清晰、正文缩进一致、表格信息可读。",
        "Normal",
        "宋体",
        12,
        False,
        first_line_cm=0.74,
    )

    table = document.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    headers = ["项目", "标准", "说明"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        set_run_font(run, "宋体", 10.5, bold=True)
    rows = [
        ("标题", "黑体 16pt 加粗", "一级标题进入目录"),
        ("正文", "宋体 12pt 首行缩进", "正文保持 1.5 倍行距"),
    ]
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            cell = table.rows[row_index].cells[col_index]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            set_run_font(run, "宋体", 10.5, bold=False)

    document.core_properties.author = "Codex"
    document.core_properties.title = "official-report fixture"
    document.save(DOCX_PATH)
    normalize_docx_zip(DOCX_PATH)


def normalize_docx_zip(path: Path) -> None:
    """固定 DOCX ZIP 时间戳，保证 fixture 哈希可复现。"""
    original_items: list[tuple[zipfile.ZipInfo, bytes]] = []
    with zipfile.ZipFile(path, "r") as source:
        for item in source.infolist():
            normalized = zipfile.ZipInfo(item.filename, FIXED_ZIP_DATE)
            normalized.compress_type = item.compress_type
            normalized.external_attr = item.external_attr
            original_items.append((normalized, source.read(item.filename)))

    with zipfile.ZipFile(path, "w") as target:
        for item, data in original_items:
            target.writestr(item, data)


def build_snapshot() -> dict:
    """生成与样例文档对应的最小事实快照。"""
    digest = hashlib.sha256(DOCX_PATH.read_bytes()).hexdigest()
    return {
        "schema_version": "1.0.0",
        "source_docx": "fixtures/standard/official-report.docx",
        "document_hash": f"sha256:{digest}",
        "created_at": FIXED_CREATED_AT,
        "paragraph_count": 5,
        "table_count": 1,
        "section_count": 1,
        "has_toc_field": False,
        "available_style_ids": ["Normal", "Title", "Heading1", "Heading2", "TableGrid"],
        "paragraphs": [
            {
                "element_id": "p-00001",
                "paragraph_index": 1,
                "text": "正式材料格式标准样例",
                "text_preview": "正式材料格式标准样例",
                "style_id": "Title",
                "paragraph_format": {"alignment": "center", "space_after_pt": 18},
                "run_format": {"font_east_asia": "黑体", "font_size_pt": 22, "bold": True},
                "numbering_pattern": "none",
            },
            {
                "element_id": "p-00002",
                "paragraph_index": 2,
                "text": "一、总体要求",
                "text_preview": "一、总体要求",
                "style_id": "Heading1",
                "paragraph_format": {"line_spacing_multiple": 1.5, "space_after_pt": 6},
                "run_format": {"font_east_asia": "黑体", "font_size_pt": 16, "bold": True},
                "numbering_pattern": "chinese-heading",
            },
            {
                "element_id": "p-00003",
                "paragraph_index": 3,
                "text": "本样例用于验证 format-helper v3 的语义规则草案与规则摘要生成链路。",
                "text_preview": "本样例用于验证 format-helper v3 的语义规则草案与规则摘要生成链路。",
                "style_id": "Normal",
                "paragraph_format": {
                    "first_line_indent_cm": 0.74,
                    "line_spacing_multiple": 1.5,
                    "space_after_pt": 6,
                },
                "run_format": {"font_east_asia": "宋体", "font_size_pt": 12, "bold": False},
                "numbering_pattern": "none",
            },
            {
                "element_id": "p-00004",
                "paragraph_index": 4,
                "text": "（一）材料范围",
                "text_preview": "（一）材料范围",
                "style_id": "Heading2",
                "paragraph_format": {"line_spacing_multiple": 1.5, "space_after_pt": 6},
                "run_format": {"font_east_asia": "楷体", "font_size_pt": 14, "bold": True},
                "numbering_pattern": "chinese-subheading",
            },
            {
                "element_id": "p-00005",
                "paragraph_index": 5,
                "text": "标准材料应保持标题层级清晰、正文缩进一致、表格信息可读。",
                "text_preview": "标准材料应保持标题层级清晰、正文缩进一致、表格信息可读。",
                "style_id": "Normal",
                "paragraph_format": {
                    "first_line_indent_cm": 0.74,
                    "line_spacing_multiple": 1.5,
                    "space_after_pt": 6,
                },
                "run_format": {"font_east_asia": "宋体", "font_size_pt": 12, "bold": False},
                "numbering_pattern": "none",
            },
        ],
        "tables": [
            {
                "element_id": "table-0001",
                "row_count": 3,
                "column_count": 3,
                "style_id": "TableGrid",
                "header": {
                    "font_east_asia": "宋体",
                    "font_size_pt": 10.5,
                    "bold": True,
                    "alignment": "center",
                },
                "body": {"font_east_asia": "宋体", "font_size_pt": 10.5, "bold": False},
            }
        ],
        "sections": [
            {
                "section_index": 1,
                "page_orientation": "portrait",
                "margins_cm": {"top": 2.6, "bottom": 2.4, "left": 2.8, "right": 2.6},
            }
        ],
        "toc_fields": [],
        "numbering_summary": {"has_numbering": True, "patterns": ["chinese-heading", "chinese-subheading"]},
    }


def main() -> None:
    create_docx()
    snapshot = build_snapshot()
    SNAPSHOT_PATH.parent.mkdir(parents=True, exist_ok=True)
    SNAPSHOT_PATH.write_text(json.dumps(snapshot, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(DOCX_PATH)
    print(SNAPSHOT_PATH)


if __name__ == "__main__":
    main()
